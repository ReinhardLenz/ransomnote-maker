from docx import Document
from docx.shared import Pt
import random

fonts = [
    "Aharoni", "Baguet Script", "Britannica Bold", "Courier New", "Daytona",
    "Elephant", "Forte Forward", "Georgia Pro Black", "Haettenschweiler", "Grotesque"
]

sizes = [8, 10, 12, 14]

doc = Document("example.docx")

for paragraph in doc.paragraphs:
    new_runs = []
    for run in paragraph.runs:
        for char in run.text:
            new_run = paragraph.add_run(char)
            new_run.font.name = random.choice(fonts)
            new_run.font.size = random.choice([Pt(s) for s in sizes])
        run.text = ""  # efface l'ancien run

doc.save("example-modified.docx")
